import streamlit as st
from datetime import date
from docx import Document
import os
import base64

# Configuração do tema
if 'theme' not in st.session_state:
    st.session_state.theme = 'dark'

def toggle_theme():
    if st.session_state.theme == 'light':
        st.session_state.theme = 'dark'
    else:
        st.session_state.theme = 'light'

# Função para gerar o documento .docx
def generate_docx(data):
    filename = f"{data['nome_paciente'].replace(' ', '_').lower()}_{data['data_relatorio']}.docx"
    doc = Document()
    
    # Inserir dados da empresa
    doc.add_heading('Despertar Atendimento Terapêutico', level=1)
    doc.add_paragraph("www.grupodespertar.org")
    doc.add_paragraph("FONES: (11) 2614-1375 | Cel. Wpp 11 96734-0051")
    doc.add_paragraph("RUA DR. ARGEMIRO COUTO DE BARROS, 06")
    doc.add_paragraph("CHÁCARA INGLESA, SÃO PAULO")
    
    # Dados do paciente
    doc.add_heading('Dados do Paciente:', level=2)
    doc.add_paragraph(f"Nome do Paciente: {data['nome_paciente']}")
    doc.add_paragraph(f"Idade: {data['idade']} anos")
    doc.add_paragraph(f"Ciclo Escolar: {data['ciclo_escolar']} / Turma: {data['turma']}")
    doc.add_paragraph(f"Instituição de Ensino: {data['instituicao_ensino']}")
    doc.add_paragraph(f"Data do Relatório: {data['data_relatorio']}")
    doc.add_paragraph(f"Profissional Responsável: {data['profissional_responsavel']}")

    # Evolução do Paciente
    doc.add_heading('Evolução do Paciente:', level=2)
    for comportamento, descricao in data['comportamentos'].items():
        if descricao:  # Só adiciona se houver descrição (não for "Não Aplicável")
            doc.add_paragraph(f"{comportamento}: {descricao}")

    # Intervenções
    doc.add_heading('Intervenções Mediante as Barreiras Enfrentadas:', level=2)
    doc.add_paragraph(data['intervencoes'])

    # Observações
    doc.add_heading('Observações do Profissional:', level=2)
    doc.add_paragraph(data['observacoes_profissional'])

    doc.save(filename)
    return filename

# Configuração da página
st.set_page_config(
    page_title="Clínica Despertar - Relatório Mensal ATs",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Função para criar um link de download
def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Download {file_label}</a>'
    return href

# Sidebar
st.sidebar.title("Sobre o Gerador de Relatório")
st.sidebar.info("""
O Gerador de Relatório Mensal ATs é uma ferramenta criada pela AperData para a Clínica Despertar, com o objetivo de facilitar a elaboração de relatórios psicológicos detalhados e personalizados, oferecendo aos profissionais um suporte prático e eficiente para a documentação dos atendimentos..
""")

st.sidebar.title("Entre em Contato")
st.sidebar.markdown("""
Para soluções de IA sob medida ou suporte:

- 🌐 [aperdata.com](https://aperdata.com)
- 📱 WhatsApp: **11 98854-3437**
- 📧 Email: **gabriel@aperdata.com**
""")
# Função principal da aplicação Streamlit
def main():
    st.title('Clínica Despertar - Relatório Mensal ATs')
    
    # Opções para ciclo escolar e turma
    opcoes_ciclo_escolar = [
        "Creche 0 a 03 anos", "Pré-escola 04 e 05 anos",
        "1º ano Ensino Fundamental", "2º ano Ensino Fundamental", "3º ano Ensino Fundamental",
        "4º ano Ensino Fundamental", "5º ano Ensino Fundamental", "6º ano Ensino Fundamental",
        "7º ano Ensino Fundamental", "8º ano Ensino Fundamental", "9º ano Ensino Fundamental",
        "1º ano Ensino Médio", "2º ano Ensino Médio", "3º ano Ensino Médio"
    ]
    opcoes_turma = ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J", "K", "L", "M"]
    
    # Formulário para entrada de dados
    with st.form("form"):
        st.header('Dados do Relatório')
        col1, col2 = st.columns(2)
        with col1:
            nome_paciente = st.text_input('Nome do Paciente')
            idade = st.number_input('Idade', min_value=0, max_value=150)
            ciclo_escolar = st.selectbox('Ciclo Escolar', options=opcoes_ciclo_escolar)
        with col2:
            turma = st.selectbox('Turma', options=opcoes_turma)
            instituicao_ensino = st.text_input('Instituição de Ensino')
            data_relatorio = st.date_input('Data do Relatório', value=date.today())
        profissional_responsavel = st.text_input('Nome do(a) Profissional Responsável')

        st.header('Evolução do Paciente')
        
        comportamentos = {
            "Manutenção do Contato Visual": [
                "tem mantido contato visual durante interações sociais e atividades de grupo, o que indica um avanço significativo na socialização.",
                "não tem mantido contato visual durante interações sociais e atividades de grupo, o que pode dificultar a interação social."
            ],
            "Seguir Instruções": [
                "segue instruções simples de maneira independente, demonstrando compreensão e autonomia.",
                "não segue instruções simples sem assistência, mostrando dificuldade em compreender ou seguir orientações."
            ],
           "Participação em Atividades em Grupo": [
                "participa ativamente das atividades em grupo, como aulas de judô, música e rodas de conversa, mostrando envolvimento e interesse.",
                "evita participar de atividades em grupo, preferindo o isolamento."
            ],
            "Comunicação Verbal": [
                "faz pedidos e expressa suas necessidades verbalmente, embora ainda com auxílio ocasional.",
                "não faz pedidos ou expressa suas necessidades verbalmente, necessitando de constante mediação."
            ],
            "Esperar a Vez": [
                "mostra paciência ao esperar sua vez durante atividades, melhorando sua capacidade de autocontrole.",
                "não consegue esperar sua vez durante atividades, mostrando impaciência e falta de autocontrole."
            ],
            "Interesse por Novas Habilidades": [
                "demonstra curiosidade e vontade de aprender novas habilidades, indicando motivação para o aprendizado.",
                "não demonstra interesse em aprender novas habilidades, o que pode afetar seu desenvolvimento acadêmico e social."
            ],
            "Expressão Emocional Adequada": [
                "expressa emoções de maneira apropriada, o que facilita a interação com colegas e professores.",
                "expressa emoções de maneira inapropriada, o que pode causar conflitos com colegas."
            ],
            "Interação Amigável": [
                "interage de forma amigável com colegas, promovendo um ambiente social positivo.",
                "age de forma agressiva ou não amigável com colegas, prejudicando a interação social."
            ],
            "Conclusão de Tarefas Acadêmicas": [
                "completa tarefas acadêmicas com pouca assistência, evidenciando progresso na independência.",
                "necessita de assistência constante para completar tarefas, evidenciando falta de autonomia."
            ],
            "Comportamento Sentado": [
                "mantém-se sentado durante atividades de mesa, facilitando a concentração e participação.",
                "levanta-se frequentemente durante atividades de mesa, indicando dificuldade em manter a atenção."
            ],
            "Aceitação de Redirecionamentos": [
                "aceita redirecionamentos sem resistência, mostrando flexibilidade e adaptação.",
                "resiste a redirecionamentos e orientações, mostrando inflexibilidade."
            ],
            "Compartilhamento de Brinquedos": [
                "compartilha brinquedos e materiais, favorecendo a interação social.",
                "tem dificuldade em compartilhar brinquedos e materiais, afetando a interação social."
            ],
            "Iniciativa em Conversas": [
                "inicia conversas com colegas e adultos, promovendo a comunicação espontânea.",
                "evita iniciar conversas com colegas e adultos, o que pode limitar seu desenvolvimento comunicativo."
            ],
            "Participação em Atividades Físicas": [
                "participa de atividades físicas sem resistência, o que é crucial para o desenvolvimento motor e social.",
                "resiste a participar de atividades físicas, o que é prejudicial para seu desenvolvimento motor."
            ],
            "Linguagem Apropriada": [
                "utiliza linguagem apropriada para sua idade, facilitando a comunicação e entendimento.",
                "usa linguagem inadequada para a idade, o que pode causar mal-entendidos."
            ],
            "Empatia": [
                "demonstra empatia pelos sentimentos dos outros, o que é fundamental para o desenvolvimento de relacionamentos saudáveis.",
                "não demonstra empatia pelos sentimentos dos outros, afetando negativamente suas relações."
            ],
            "Seguir Rotinas": [
                "segue rotinas diárias sem protesto, mostrando adaptação e compreensão das mesmas.",
                "protesta contra rotinas diárias, indicando resistência e dificuldade de adaptação."
            ],
            "Cumprimento de Regras": [
                "cumpre regras estabelecidas com consistência, o que promove um ambiente escolar seguro e organizado.",
                "tem dificuldade em cumprir regras estabelecidas, prejudicando o ambiente escolar."
            ]
        }

        comportamentos_selecionados = {}
        for comportamento, opcoes in comportamentos.items():
            escolha = st.selectbox(comportamento, ["Positivo", "Negativo", "Não Aplicável"])
            if escolha == "Positivo":
                comportamentos_selecionados[comportamento] = opcoes[0]
            elif escolha == "Negativo":
                comportamentos_selecionados[comportamento] = opcoes[1]
            else:
                comportamentos_selecionados[comportamento] = ""  # Não adiciona nada se for "Não Aplicável"

        st.header('Intervenções Mediante as Barreiras Enfrentadas')
        intervencoes = st.text_area('Intervenções Mediante as Barreiras Enfrentadas')

        st.header('Observações do Profissional')
        observacoes_profissional = st.text_area('Observações do Profissional')

        # Botão para gerar o relatório
        submitted = st.form_submit_button('Gerar Relatório')
        
        if submitted:
            data = {
                'nome_paciente': nome_paciente,
                'idade': idade,
                'ciclo_escolar': ciclo_escolar,
                'turma': turma,
                'instituicao_ensino': instituicao_ensino,
                'data_relatorio': data_relatorio.strftime('%d-%m-%Y'),
                'profissional_responsavel': profissional_responsavel,
                'comportamentos': comportamentos_selecionados,
                'intervencoes': intervencoes,
                'observacoes_profissional': observacoes_profissional
            }

            filename = generate_docx(data)
            st.success('Relatório gerado com sucesso!')
            st.markdown(get_binary_file_downloader_html(filename, 'Relatório'), unsafe_allow_html=True)

if __name__ == '__main__':
    main()
